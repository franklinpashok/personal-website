import pandas as pd
from docx import Document

# --- 1. GENERATE DUMMY EXCEL (Cloud Budget) ---
print("Generating 'cloud_budget_2026.xlsx'...")

data = {
    "Service": ["AWS EKS Cluster", "RDS (PostgreSQL)", "Nat Gateway", "CloudFront", "Lambda Functions"],
    "Region": ["eu-south-2 (Spain)", "eu-south-2 (Spain)", "us-east-1", "Global", "eu-south-2 (Spain)"],
    "Monthly_Cost_USD": [144.00, 85.50, 32.00, 15.00, 4.20],
    "Status": ["Active", "Active", "Deprecated", "Active", "Active"],
    "Owner": ["DevOps Team", "Data Team", "Legacy System", "Frontend Team", "Backend Team"]
}

df = pd.DataFrame(data)
df.to_excel("cloud_budget_2026.xlsx", index=False)
print("✅ Excel file created.")

# --- 2. GENERATE DUMMY WORD DOC (Incident Report) ---
print("Generating 'incident_report_feb_2026.docx'...")

doc = Document()
doc.add_heading('Security Incident Report - #SEC-2026-001', 0)

doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'On February 4, 2026, the monitoring system detected an unusual spike in API requests '
    'originating from a single IP address in North Korea. The automated WAF rules successfully '
    'blocked the traffic after 45 seconds.'
)

doc.add_heading('2. Impact Analysis', level=1)
doc.add_paragraph(
    'No user data was compromised. However, the "Login" service experienced higher latency '
    '(approx. 400ms) during the attack window.'
)

doc.add_heading('3. Required Actions', level=1)
p = doc.add_paragraph()
p.add_run('Immediate: ').bold = True
p.add_run('Rotate all IAM credentials for the "CI/CD Pipeline" role.')

p2 = doc.add_paragraph()
p2.add_run('Long-term: ').bold = True
p2.add_run('Implement Geo-Blocking at the CloudFront level for non-business regions.')

doc.save("incident_report_feb_2026.docx")
print("✅ Word document created.")